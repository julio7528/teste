import sys
import os
import pandas as pd
import logging
import zipfile
from xml.etree import ElementTree as ET
from docx import Document
import win32com.client
import re
current_dir = os.path.dirname(os.path.abspath(__file__))
src_dir = os.path.join(current_dir, '..') 
sys.path.append(src_dir)

from config.config import load_config, get_caminho_de_para 

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

load_config()

def read_excel_file(caminhoArquivo, header=2, sheet_name=0):

    try:
        if not os.path.exists(caminhoArquivo):
            raise FileNotFoundError(f"O arquivo '{caminhoArquivo}' não foi encontrado.")
        
        df = pd.read_excel(caminhoArquivo, header=header, sheet_name=sheet_name)

        if df.empty:
            logging.warning("O arquivo está vazio ou não contém dados suficientes.")

        return df

    except Exception as e:
        logging.error(f"Erro ao tentar ler o arquivo: {e}")
        return None
    

def capture_code_from_docx(caminho_arquivo, prefixos=["Código:", "CÓDIGO DO MATERIAL", "CÓDIGO", "Código do Produto: ", "CÓDIGO INTERNO DO MATERIAL"]):
    try:
        # Ordena os prefixos por comprimento decrescente para dar prioridade aos mais específicos
        prefixos = sorted(prefixos, key=len, reverse=True)
        
        # Abre o arquivo Word
        documento = Document(caminho_arquivo)
        
        # Verifica parágrafos no documento
        for i, paragrafo in enumerate(documento.paragraphs):
            texto = paragrafo.text.strip()
            for prefixo in prefixos:  # Itera pelos prefixos em ordem de prioridade
                if texto.startswith(prefixo):
                    # Captura o código após o prefixo
                    codigo = texto[len(prefixo):].strip()
                    
                    if not codigo or codigo == ":":
                        for j in range(i + 1, len(documento.paragraphs)):
                            prox_paragrafo = documento.paragraphs[j].text.strip()
                            if prox_paragrafo:
                                codigo = prox_paragrafo
                                break
                    
                    # Retorna se encontrar um código válido
                    if codigo and codigo != ":":
                        return codigo.replace('.', '')
        
        # Verifica tabelas no documento se o código não foi encontrado nos parágrafos
        for tabela in documento.tables:
            for i, linha in enumerate(tabela.rows):
                celulas = [celula.text.strip() for celula in linha.cells]
                for j, texto in enumerate(celulas):
                    for prefixo in prefixos:
                        if texto.startswith(prefixo):
                            codigo = texto[len(prefixo):].strip()
                            
                            if not codigo and j + 1 < len(celulas):
                                codigo = celulas[j + 1].strip()
                            
                            if codigo and codigo != ":":
                                return codigo.replace('.', '')

        # Se nenhum código for encontrado
        return "Nenhum código encontrado."
    except Exception as e:
        logging.error(f"Erro ao tentar capturar o código do documento: {e}")
        return "Erro ao capturar o código."

def get_headers_texts(caminho_arquivo):
    header_texts_data = []
    documento = Document(caminho_arquivo)

    for secao in documento.sections:
        header = secao.header
        for tabela in header.tables:
            for linha in tabela.rows:
                celulas = [celula.text.strip() for celula in linha.cells]
                for texto in celulas:
                    header_texts_data.append(texto)

    return header_texts_data

def capture_type_from_headers(caminho_arquivo):
    header_texts_data = " ".join(get_headers_texts(caminho_arquivo)).replace("\n", " ").strip()
    if "FICHA DE ANÁLISE" in header_texts_data:
        return "FICHA DE ANÁLISE"
    elif "INSTRUÇÃO DE EMBALAGEM" in header_texts_data:
        return "INSTRUÇÃO DE EMBALAGEM"
    elif "INSTRUÇÃO DE FABRICAÇÃO" in header_texts_data:
        return "INSTRUÇÃO DE FABRICAÇÃO"
    elif "ESPECIFICAÇÃO DE MATERIAL DE EMBALAGEM" in header_texts_data:
        return "EME"
    elif "ESPECIFICAÇÃO DE MATÉRIA PRIMA" in header_texts_data:
        return "EMP"
    elif "ESPECIFICAÇÃO DE PRODUTO ACABADO" in header_texts_data:
        return "EPA"
    elif "ESPECIFICAÇÃO DE PRODUTO EM ESTABILIDADE" in header_texts_data:
        return "EPE"
    elif "ESPECIFICAÇÃO DE PRODUTO INTERMEDIÁRIO" in header_texts_data:
        return "EPI"
    elif "MÉTODO DE ANÁLISE DE MATERIAL DE EMBALAGEM" in header_texts_data:
        return "MAME"
    elif "MÉTODO DE ANÁLISE DE MATÉRIA PRIMA" in header_texts_data:
        return "MAMP"
    elif "MÉTODO DE ANÁLISE DE PRODUTO ACABADO" in header_texts_data:
        return "MAPA"
    elif "METODOLOGIA DE ANÁLISE DE PRODUTO EM ESTABILIDADE" in header_texts_data:
        return "MAPE"
    elif "MÉTODO DE ANÁLISE PRODUTO DE INTERMEDIÁRIO" in header_texts_data:
        return "MAPI"
    else:
        return None

def capture_code_from_headers(caminho_arquivo, prefixos=["Código do Produto: "]):
    """
    Função independente para buscar códigos e o texto 'INSTRUÇÃO DE EMBALAGEM' nos cabeçalhos do documento.
    Verifica parágrafos e tabelas dentro dos cabeçalhos de cada seção.

    Args:
        caminho_arquivo (str): Caminho para o arquivo .docx.
        prefixos (list): Lista de prefixos para identificar os códigos.

    Returns:
        tuple: (código encontrado, instrução encontrada), ou ("Nenhum código encontrado", "Nenhuma instrução encontrada").
    """
    documento = Document(caminho_arquivo)

    # Variáveis para armazenar os resultados
    codigo_encontrado = None
    instrucao_encontrada = None

    # Itera pelas seções do documento para acessar os cabeçalhos
    for i, secao in enumerate(documento.sections):
        header = secao.header  # Obtém o cabeçalho da seção

        # Verifica parágrafos no cabeçalho
        for paragrafo in header.paragraphs:
            texto = paragrafo.text.strip()

            # Verifica prefixos para código
            for prefixo in prefixos:
                if texto.startswith(prefixo):
                    codigo = texto[len(prefixo):].strip()
                    if codigo and codigo != ":":
                        codigo_encontrado = codigo.replace('.', '')

            # Verifica "INSTRUÇÃO DE EMBALAGEM"
            if "INSTRUÇÃO DE EMBALAGEM" in texto.upper():
                instrucao_encontrada = texto.strip()

            if "INSTRUÇÃO DE FABRICAÇÃO" in texto.upper():
                instrucao_encontrada = texto.strip()
            

        # Verifica tabelas no cabeçalho
        for t, tabela in enumerate(header.tables):
            for linha in tabela.rows:
                celulas = [celula.text.strip() for celula in linha.cells]
                for j, texto in enumerate(celulas):
                    # Verifica prefixos para código
                    for prefixo in prefixos:
                        if texto.startswith(prefixo):
                            codigo = texto[len(prefixo):].strip()
                            if not codigo and j + 1 < len(celulas):
                                codigo = celulas[j + 1].strip()
                            if codigo and codigo != ":":
                                codigo_encontrado = codigo.replace('.', '')

                    # Verifica "INSTRUÇÃO DE EMBALAGEM"
                    if "INSTRUÇÃO DE EMBALAGEM" in texto.upper():
                        instrucao_encontrada = texto.strip()

                    if "INSTRUÇÃO DE FABRICAÇÃO" in texto.upper():
                        instrucao_encontrada = texto.strip()

    # Valores padrão caso não sejam encontrados
    if not codigo_encontrado:
        codigo_encontrado = "Nenhum código encontrado"
    if not instrucao_encontrada:
        instrucao_encontrada = "Nenhuma instrução encontrada"

    return codigo_encontrado, instrucao_encontrada

def convert_doc_to_docx(input_path, output_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=16)  # 16 corresponde ao formato .docx
        doc.Close()
        word.Quit()
    except Exception as e:
        logging.error(f"Erro ao tentar converter o arquivo doc_to_docx: {e}")


def convert_docx_to_doc(input_path, output_path):
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=0)
        doc.Close()
        word.Quit()
    except Exception as e:
        logging.error(f"Erro ao tentar converter o arquivo docx_to_doc: {e}")



def write_document_content_to_file(caminho_arquivo, caminho_saida):
    """
    Escreve o conteúdo de um documento Word (.docx) em um arquivo de texto, incluindo cabeçalhos, parágrafos e tabelas.

    Args:
        caminho_arquivo (str): Caminho para o arquivo .docx.
        caminho_saida (str): Caminho para o arquivo de saída (.txt).
    """
    # Abre o documento
    documento = Document(caminho_arquivo)

    with open(caminho_saida, "w", encoding="utf-8") as arquivo_saida:
        arquivo_saida.write("=== Conteúdo do Documento ===\n\n")

        # Escreve os cabeçalhos do documento (parágrafos e tabelas)
        arquivo_saida.write("Cabeçalhos:\n\n")
        for i, section in enumerate(documento.sections):
            header = section.header
            arquivo_saida.write(f"--- Cabeçalho da Seção {i + 1} ---\n")
            
            # Escreve os parágrafos no cabeçalho
            for paragrafo in header.paragraphs:
                texto = paragrafo.text.strip()
                if texto:
                    arquivo_saida.write(f"{texto}\n")

            # Escreve as tabelas no cabeçalho
            if header.tables:
                arquivo_saida.write("\nConteúdo das Tabelas no Cabeçalho:\n")
                for t, tabela in enumerate(header.tables):
                    arquivo_saida.write(f"Tabela {t + 1}:\n")
                    for linha in tabela.rows:
                        linha_texto = [celula.text.strip() for celula in linha.cells]
                        arquivo_saida.write("\t".join(linha_texto) + "\n")
                    arquivo_saida.write("-" * 40 + "\n")
            
            arquivo_saida.write("\n")

        # Escreve os parágrafos do documento
        arquivo_saida.write("Parágrafos:\n\n")
        for i, paragrafo in enumerate(documento.paragraphs):
            texto = paragrafo.text.strip()
            if texto:
                arquivo_saida.write(f"Parágrafo {i + 1}: {texto}\n")

        arquivo_saida.write("\n=== Conteúdo das Tabelas ===\n\n")

        # Escreve o conteúdo de cada tabela no corpo do documento
        for t, tabela in enumerate(documento.tables):
            arquivo_saida.write(f"Tabela {t + 1}:\n")
            for linha in tabela.rows:
                linha_texto = [celula.text.strip() for celula in linha.cells]
                arquivo_saida.write("\t".join(linha_texto) + "\n")  # Exibe as células da linha separadas por tabulação
            arquivo_saida.write("-" * 40 + "\n")  # Separador entre tabelas

        arquivo_saida.write("\n=== Fim do Documento ===")

    print(f"Conteúdo salvo em: {caminho_saida}")



def captura_tabela_embalagem(arq, texto_inicial_tabela):
    # Carregar o documento Word
    documento = Document(arq)

    # Lista para armazenar os dados extraídos
    dados_extracao = []

    # Iterar pelas tabelas no documento
    for tabela in documento.tables:
        # Verificar se a primeira célula da tabela contém o texto inicial desejado
        if tabela.rows[0].cells[0].text.strip().startswith(texto_inicial_tabela):
            for linha in tabela.rows[1:]:  # Ignorar o cabeçalho
                # Verificar se há pelo menos 5 colunas na linha (para evitar erros de índice)
                if len(linha.cells) >= 5:
                    try:
                        # Capturar o código (procurar na posição 0 ou 1)
                        codigo = linha.cells[0].text.strip() if linha.cells[0].text.strip().isdigit() else linha.cells[1].text.strip()

                        # Capturar os demais valores
                        descricao = linha.cells[2].text.strip()  # Descrição
                        quantidade = linha.cells[3].text.strip()  # Quantidade
                        unidade = linha.cells[4].text.strip()  # Unidade

                        # Verificar se os dados são válidos
                        if codigo.isdigit() and quantidade.replace('.', '', 1).isdigit():
                            dados_extracao.append({
                                "Código": codigo,
                                "Descrição": descricao,
                                "Quantidade": quantidade,
                                "Unidade": unidade
                            })
                    except IndexError:
                        # Ignorar linhas com estrutura inesperada
                        continue
            break  # Processar apenas a primeira tabela encontrada com o texto inicial

    return dados_extracao


def captura_tabela_fabricacao(arq, texto_inicial_tabela):
    # Carregar o documento Word
    documento = Document(arq)

    # Lista para armazenar os dados extraídos
    dados_extracao = []

    # Iterar pelas tabelas no documento
    for tabela in documento.tables:
        # Verificar se a primeira célula da tabela contém o texto inicial desejado
        if tabela.rows[0].cells[0].text.strip().startswith(texto_inicial_tabela):
            for linha in tabela.rows[1:]:  # Ignorar o cabeçalho
                # Verificar se há pelo menos 5 colunas na linha (para evitar erros de índice)
                if len(linha.cells) >= 5:
                    try:
                        # Capturar o código (procurar na posição 0 ou 1)
                        codigo = linha.cells[0].text.strip() if linha.cells[0].text.strip().isdigit() else linha.cells[1].text.strip()

                        # Capturar os demais valores
                        descricao = linha.cells[2].text.strip()  # Descrição
                        dcb = linha.cells[3].text.strip()  # DCB
                        QuantTeórica  = linha.cells[4].text.strip()  # Unidade
                        formula = linha.cells[5].text.strip()
                        formulaUni = linha.cells[6].text.strip()
                        # Verificar se os dados são válidos
                        if codigo.isdigit() and dcb.replace('.', '', 1).isdigit():
                            dados_extracao.append({
                                "Código": codigo,
                                "Descrição": descricao,
                                "DCB": dcb,
                                "QuantTeórica": QuantTeórica,
                                "Fórmula":formula,
                                "Fórmula Unitária":formulaUni
                            })
                    except IndexError:
                        # Ignorar linhas com estrutura inesperada
                        continue
            break  # Processar apenas a primeira tabela encontrada com o texto inicial

    return dados_extracao


def captura_codigo_nucleo(file_path, keyword):

    pattern = r"\b\d{8}\b"
    # Abrir o arquivo Word
    doc = Document(file_path)

    # Lista para armazenar códigos encontrados
    matched_codes = []

    # Procurar a palavra-chave nos parágrafos
    for paragraph in doc.paragraphs:
        if keyword in paragraph.text:
            codes = re.findall(pattern, paragraph.text)
            matched_codes.extend(codes)

    # Procurar a palavra-chave nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if keyword in cell.text:
                    codes = re.findall(pattern, cell.text)
                    matched_codes.extend(codes)

    # Remover duplicatas
    return list(set(matched_codes))




# arq = r"C:\RPA\RPA001_Garantia_De_Qualidade\data\IF-IE\OP-0017.docx"
# texto_inicial_tabela = "Componentes – Núcleo"
# df_embalagem = captura_tabela_fabricacao(arq, texto_inicial_tabela)

# for item in df_embalagem:
#     print(f"Código: {item['Código']}")
#     print(f"Descrição: {item['Descrição']}")
#     print(f"DCB: {item['DCB']}")
#     print(f"QuantTeórica: {item['QuantTeórica']}")
#     print(f"Fórmula: {item['Fórmula']}")
#     print(f"Fórmula Unitária: {item['Fórmula Unitária']}")




# arq = r"C:\RPA\RPA001_Garantia_De_Qualidade\data\IF-IE\OP-0017.docx"
# arqText = r"C:\RPA\RPA001_Garantia_De_Qualidade\data\IF-IE\OP-0017.txt"
# write_document_content_to_file(arq, arqText)

