import sys
import os
import pandas as pd
import logging
from xml.etree import ElementTree as ET
from docx import Document
import win32com.client
import re
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy

current_dir = os.path.dirname(os.path.abspath(__file__))
src_dir = os.path.join(current_dir, '..') 
sys.path.append(src_dir)

from config.config import load_config, get_caminho_de_para 

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

load_config()


def edit_file_ficha(caminho_arquivo, texto_procurado, novo_texto, caminho_arquivo_salvo):

    # Abre o arquivo Word
    documento = Document(caminho_arquivo)

    
    # Percorre as tabelas do documento
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if texto_procurado in celula.text.strip():  # Verifica se o texto corresponde ao procurado
                  
                    celula.text = celula.text.replace(texto_procurado, novo_texto)  # Altera o texto

                    # Ajusta a formatação da célula
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            run.font.name = "Calibri"
                            run.font.size = Pt(12)
                            run.bold = True

                            # Garante que a fonte seja Calibri (Corpo)
                            rFonts = run._element.rPr.rFonts
                            rFonts.set(qn("w:ascii"), "Calibri")
                            rFonts.set(qn("w:hAnsi"), "Calibri")
                    


    for i, secao in enumerate(documento.sections):
        rodape = secao.footer
      

        for paragrafo in rodape.paragraphs:
           
            
            # Localiza a revisão com regex
            match = re.search(r"Revisão (\d+)", paragrafo.text)
            if match:
                numero_revisao = int(match.group(1))  # Captura o número da revisão
                nova_revisao = f"Revisão {numero_revisao + 1:02d}"  # Incrementa +1
                
                # Substitui o texto no parágrafo
                paragrafo.text = re.sub(r"Revisão \d+", nova_revisao, paragrafo.text)
               

    # Salva o documento com as alterações
    documento.save(caminho_arquivo_salvo)


def edit_file_eme(caminho_arquivo, texto_procurado, novo_texto, caminho_arquivo_salvo):
    # Abre o arquivo Word
    documento = Document(caminho_arquivo)

    # 1. Substitui o texto_procurado nas tabelas
    texto_alterado = False
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                if texto_procurado in celula.text.strip():
                    
                    celula.text = celula.text.replace(texto_procurado, novo_texto)


                    # Aplica a formatação Arial e tamanho 10
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            run.font.name = "Arial"  # Define a fonte Arial
                            run.font.size = Pt(10)  # Define o tamanho da fonte 10
                    
                    texto_alterado = True

    if not texto_alterado:
        print("Texto procurado não encontrado no documento.")


    # 2. Localiza a tabela de revisão e insere uma nova linha no topo
    tabela_revisao_encontrada = False
    for tabela in documento.tables:
        if "REVISÃO" in tabela.cell(0, 0).text.strip():
            tabela_revisao_encontrada = True

            # Determina o maior número de revisão
            maior_revisao = 0
            for linha in tabela.rows[1:]:  # Ignora o cabeçalho
                try:
                    revisao_atual = int(re.search(r"\d+", linha.cells[0].text.strip()).group())
                    maior_revisao = max(maior_revisao, revisao_atual)
                except (ValueError, AttributeError):
                    continue

            # Cria a nova revisão
            nova_revisao = maior_revisao + 1
            data_atual = datetime.now().strftime("%d/%m/%Y")

            # Adiciona uma linha no final para evitar sobrescrever linhas existentes
            tabela.add_row()
            for i in range(len(tabela.rows) - 1, 1, -1):  # Move as linhas para baixo
                for j, cell in enumerate(tabela.rows[i - 1].cells):
                    tabela.rows[i].cells[j].text = cell.text

            # Preenche a nova primeira linha com os dados
            primeira_linha = tabela.rows[1].cells
            primeira_linha[0].text = str(nova_revisao).zfill(2)
            primeira_linha[1].text = "-"  # Coluna ITEM
            primeira_linha[2].text = "Revisão dos documentos mediante ao CM-TBS-00728;"
            primeira_linha[3].text = data_atual

            # Aplica a formatação em todas as células da tabela
            for linha in tabela.rows:
                for i, cell in enumerate(linha.cells):  # Percorre as células da linha
                    # Aplica a formatação de fonte e tamanho
                    for paragrafo in cell.paragraphs:
                        for run in paragrafo.runs:
                            run.font.name = "Arial"  # Define a fonte Arial
                            run.font.size = Pt(10)  # Define o tamanho da fonte 10
                    
                    # Alinha as células específicas ao centro (primeira, segunda e quarta coluna)
                    if i == 0 or i == 1 or i == 3:  # Primeira, segunda e quarta coluna
                        for paragrafo in cell.paragraphs:
                            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

            break

    if not tabela_revisao_encontrada:
        print("Tabela de revisão não encontrada no documento.")

    # 3. Salva o documento com as alterações
    documento.save(caminho_arquivo_salvo)
 

def edit_file_mame(caminho_arquivo, texto_procurado, novo_texto, caminho_arquivo_salvo):
   
    # Abre o arquivo Word
    documento = Document(caminho_arquivo)
    #1. Substitui o texto_procurado nos parágrafos
    for paragrafo in documento.paragraphs:
        if texto_procurado in paragrafo.text:
            
            paragrafo.text = paragrafo.text.replace(texto_procurado, novo_texto)

            # Aplica a formatação (Fonte Arial e Tamanho 10) ao novo texto
            for run in paragrafo.runs:  # Itera sobre todos os 'runs' no parágrafo
                run.font.name = 'Arial'  # Define a fonte Arial
                run.font.size = Pt(10)   # Define o tamanho da fonte 10
            
            texto_alterado = True

    # Verifica se o texto foi alterado
    if not texto_alterado:
        print("Texto procurado não encontrado no documento.")

    # 2. Localiza a tabela de revisão e insere uma nova linha no topo
    tabela_revisao_encontrada = False
    for tabela in documento.tables:
        if "REVISÃO" in tabela.cell(0, 0).text.strip():
            tabela_revisao_encontrada = True

            # Determina o maior número de revisão
            maior_revisao = 0
            for linha in tabela.rows[1:]:  # Ignora o cabeçalho
                try:
                    revisao_atual = int(re.search(r"\d+", linha.cells[0].text.strip()).group())
                    maior_revisao = max(maior_revisao, revisao_atual)
                except (ValueError, AttributeError):
                    continue

            # Cria a nova revisão
            nova_revisao = maior_revisao + 1
            data_atual = datetime.now().strftime("%d/%m/%Y")

            # Adiciona uma linha no final para evitar sobrescrever linhas existentes
            tabela.add_row()
            for i in range(len(tabela.rows) - 1, 1, -1):  # Move as linhas para baixo
                for j, cell in enumerate(tabela.rows[i - 1].cells):
                    tabela.rows[i].cells[j].text = cell.text

            # Preenche a nova primeira linha com os dados
            primeira_linha = tabela.rows[1].cells
            primeira_linha[0].text = str(nova_revisao).zfill(2)
            primeira_linha[1].text = "-"  # Coluna ITEM
            primeira_linha[2].text = "Revisão dos documentos mediante ao CM-TBS-00728;"
            primeira_linha[3].text = data_atual


            # Aplica a formatação em todas as células da tabela
            for linha in tabela.rows:
                for i, cell in enumerate(linha.cells):  # Percorre as células da linha
                    # Aplica a formatação de fonte e tamanho
                    for paragrafo in cell.paragraphs:
                        for run in paragrafo.runs:
                            run.font.name = "Arial"  # Define a fonte Arial
                            run.font.size = Pt(10)  # Define o tamanho da fonte 10
                    
                    # Alinha as células específicas ao centro (primeira, segunda e quarta coluna)
                    if i == 0 or i == 1 or i == 3:  # Primeira, segunda e quarta coluna
                        for paragrafo in cell.paragraphs:
                            paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
         
            break

    if not tabela_revisao_encontrada:
        print("Tabela de revisão não encontrada no documento.")

    # 3. Salva o documento com as alterações
    documento.save(caminho_arquivo_salvo)
    

def edit_file_embalagem_fabricacao(caminho_arquivo, texto_procurado, novo_texto, caminho_arquivo_salvo):
    # Abre o arquivo Word
    documento = Document(caminho_arquivo)

    # Percorre as seções do documento para acessar os cabeçalhos
    for secao in documento.sections:
        cabecalho = secao.header  # Obtém o cabeçalho da seção

        # Percorre as tabelas dentro do cabeçalho
        for tabela in cabecalho.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    if texto_procurado in celula.text.strip():  # Verifica se o texto corresponde ao procurado
                        celula.text = celula.text.replace(texto_procurado, novo_texto)  # Altera o texto

                    # Aplica a formatação de fonte e tamanho
                    for paragrafo in celula.paragraphs:
                        for run in paragrafo.runs:
                            run.font.name = "Arial"  # Define a fonte Arial
                            run.font.size = Pt(9)  # Define o tamanho da fonte 9

        # Percorre os parágrafos no cabeçalho
        for paragrafo in cabecalho.paragraphs:
            if texto_procurado in paragrafo.text.strip():  # Verifica se o texto procurado está no parágrafo
                paragrafo.text = paragrafo.text.replace(texto_procurado, novo_texto)  # Altera o texto

            # Aplica a formatação de fonte e tamanho
            for run in paragrafo.runs:
                run.font.name = "Arial"  # Define a fonte Arial
                run.font.size = Pt(9)  # Define o tamanho da fonte 9



    tabela_revisao_encontrada = False
    for tabela in documento.tables:
        # Imprime o conteúdo da célula de cabeçalho para verificar
        cabecalho = tabela.cell(0, 0).text.strip()
        
        if "Nº Revisão" == cabecalho:  # Verifica se a tabela é a que você deseja
            tabela_revisao_encontrada = True

            # Determina o maior número de revisão
            maior_revisao = 0
            for linha in tabela.rows[1:]:  # Ignora o cabeçalho
                try:
                    # Tenta capturar o número da revisão da primeira célula
                    texto_revisao = linha.cells[0].text.strip()
                    
                    # A expressão regular tentará pegar o número, mas só se encontrar um número válido
                    if re.search(r"\d+", texto_revisao):
                        revisao_atual = int(re.search(r"\d+", texto_revisao).group())
                        maior_revisao = max(maior_revisao, revisao_atual)
                    else:
                        print("Nenhum número de revisão encontrado nesta linha.")
                except (ValueError, AttributeError) as e:
                    print(f"Erro ao processar a linha: {e}")
                    continue

            # Cria a nova revisão
            nova_revisao = maior_revisao + 1
            data_atual = datetime.now().strftime("%d/%m/%Y")

            # Adiciona uma linha no final para evitar sobrescrever linhas existentes
            nova_linha = tabela.add_row()

            # Preenche a nova linha com os dados
            nova_linha.cells[0].text = str(nova_revisao).zfill(2)
            nova_linha.cells[1].text = "Revisão dos documentos mediante ao CM-TBS-00728;"
            nova_linha.cells[2].text = data_atual

            # Copia a formatação da célula acima para a nova célula
            for i, cell in enumerate(nova_linha.cells):
                # Acessa a célula da linha acima usando o índice
                celula_acima = tabela.cell(len(tabela.rows) - 2, i)  # Índice da linha acima (penúltima linha)

                # Copiar a formatação de texto da célula acima
                for par in celula_acima.paragraphs:
                    if par.text.strip():  # Se houver texto
                        for run in par.runs:
                            # Aplica a formatação da célula acima
                            for new_paragraph in cell.paragraphs:
                                for new_run in new_paragraph.runs:
                                    new_run.font.name = run.font.name
                                    new_run.font.size = run.font.size
                                    new_run.font.bold = run.font.bold
                                    new_run.font.italic = run.font.italic
                                    new_run.font.underline = run.font.underline

                # Copiar a cor de fundo (shading) se existir
                try:
                    celula_acima_format = celula_acima._element.xpath(".//w:shd")
                    if celula_acima_format:
                        # Adiciona o shading na nova célula sem remover da original
                        if not cell._element.xpath(".//w:shd"):  # Verifica se já existe fundo
                            shading_copy = deepcopy(celula_acima_format[0])  # Cria uma cópia do elemento
                            cell._element.get_or_add_tcPr().append(shading_copy)
                except Exception as e:
                    print(f"Erro ao copiar a formatação de fundo: {e}")

            # Aplica a formatação de fonte Arial, tamanho 9 e alinha as células
            for cell in nova_linha.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"  # Define a fonte Arial
                        run.font.size = Pt(9)  # Define o tamanho da fonte 9

                    # Alinha o texto da primeira, segunda e quarta coluna ao centro
                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            break

    if not tabela_revisao_encontrada:
        print("Tabela de revisão não encontrada no documento.")

    # Salva o documento com as alterações
    documento.save(caminho_arquivo_salvo)



def editar_celula_codigo_embalagem(caminho_arquivo, codigo_antigo, novo_codigo, caminho_arquivo_salvo):
    """
    Localiza uma célula com o código antigo em uma tabela e atualiza o conteúdo 
    de acordo com as regras especificadas:
    - Substitui o código antigo pelo novo caso ele comece com "3".
    - Caso contrário, adiciona o novo código entre parênteses ao lado do código original.
    - Alinha o texto ao centro e aplica a formatação de fonte Arial, tamanho 9.

    Args:
        caminho_arquivo (str): Caminho do arquivo Word (.docx) original.
        codigo_antigo (str): Código antigo que será localizado na tabela.
        novo_codigo (str): Novo código a ser adicionado ou substituído.
        caminho_arquivo_salvo (str): Caminho do arquivo Word (.docx) salvo.
    """
    # Abre o arquivo Word
    documento = Document(caminho_arquivo)
    codigo_encontrado = False  # Flag para verificar se o código foi localizado

    # Percorre as tabelas do documento
    for tabela in documento.tables:
        # Itera pelas linhas da tabela
        for linha in tabela.rows:
            for celula in linha.cells:
                # Localiza a célula que contém o código antigo
                if celula.text.strip() == codigo_antigo:
                    codigo_encontrado = True

                    # Regra de substituição ou adição
                    if codigo_antigo.startswith("3"):
                        # Substitui o código antigo pelo novo
                        celula.text = novo_codigo
                    else:
                        # Adiciona o novo código entre parênteses ao lado do original
                        celula.text = f"{codigo_antigo} ({novo_codigo})"

                    # Aplica a formatação de fonte e alinhamento
                    for paragrafo in celula.paragraphs:
                        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Alinha ao centro
                        for run in paragrafo.runs:
                            run.font.name = "Arial"  # Define a fonte Arial
                            run.font.size = Pt(9)  # Define o tamanho da fonte para 9pt

                    # Sai do loop após encontrar e editar o código
                    break
            if codigo_encontrado:
                break
        if codigo_encontrado:
            break

    # Verifica se o código antigo foi encontrado e editado
    if not codigo_encontrado:
        print(f"Código '{codigo_antigo}' não encontrado no documento.")

    # Salva o documento com as alterações
    documento.save(caminho_arquivo_salvo)




def adicionar_nova_linha_com_codigo_embalagem(caminho_arquivo, novo_codigo, descricao, quantidade, unidade, caminho_arquivo_salvo):
    # Abre o arquivo Word
    documento = Document(caminho_arquivo)

    # Percorre as tabelas do documento
    for tabela_index, tabela in enumerate(documento.tables):
        
        # Verifica se o cabeçalho está na primeira linha
        if "Componentes – Material de Embalagem" in tabela.rows[0].cells[0].text:
            

            # Acha a posição do cabeçalho (linha 2, índice 1)
            posicao_cabecalho = 2  # Cabeçalho está na segunda linha

            # Copia os dados da linha 3 (primeira linha de dados reais)
            linha_base = [celula.text for celula in tabela.rows[posicao_cabecalho + 1].cells]

            # Modifica o código (primeira célula) da nova linha e deixa a descrição em branco
            linha_base[0] = novo_codigo
            linha_base[1] = ""  # Deixa a descrição em branco
            linha_base[2] = descricao  # Insere a nova descrição (na coluna 3)
            linha_base[3] = quantidade  # Insere a quantidade (na coluna 4)
            linha_base[4] = unidade  # Insere a unidade (na coluna 5)

            # Insere a nova linha após o cabeçalho e antes da linha 3
            linhas_existentes = [[celula.text for celula in linha.cells] for linha in tabela.rows]
            linhas_existentes.insert(posicao_cabecalho + 1, linha_base)

            # Remove todas as linhas da tabela original
            for _ in range(len(tabela.rows)):
                tabela._element.remove(tabela.rows[0]._element)

            # Recria a tabela com a primeira linha ajustada
            for linha_index, linha_dados in enumerate(linhas_existentes):
                nova_linha = tabela.add_row()
                for idx, valor in enumerate(linha_dados):
                    nova_linha.cells[idx].text = valor

            # Exibe a tabela atualizada no console
            for linha_index, linha in enumerate(tabela.rows):
                conteudo_celulas = [celula.text.strip() for celula in linha.cells]
            break  # Processa apenas a tabela correta

    # Salva o documento com as alterações
    documento.save(caminho_arquivo_salvo)

def substituir_codigo_nucleo(file_path, codigo_antigo, codigo_novo, output_path):
    """
    Substitui o código antigo pelo novo em parágrafos e tabelas de um documento Word.
    Aplica fonte Arial e tamanho 9 aos textos substituídos.

    Args:
        file_path (str): Caminho do arquivo Word original.
        codigo_antigo (str): Código antigo a ser substituído.
        codigo_novo (str): Novo código para substituir o antigo.
        output_path (str): Caminho para salvar o documento modificado.
    """
    # Abrir o arquivo Word
    doc = Document(file_path)

    # Substituir o código nos parágrafos
    for paragraph in doc.paragraphs:
        if codigo_antigo in paragraph.text:
            paragraph.text = paragraph.text.replace(codigo_antigo, codigo_novo)
            # Aplica a formatação aos runs do parágrafo
            for run in paragraph.runs:
                run.font.name = "Arial"  # Define a fonte Arial
                run.font.size = Pt(9)  # Define o tamanho 9pt

    # Substituir o código nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if codigo_antigo in cell.text:
                    cell.text = cell.text.replace(codigo_antigo, codigo_novo)
                    # Aplica a formatação aos runs dos parágrafos na célula
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Arial"  # Define a fonte Arial
                            run.font.size = Pt(9)  # Define o tamanho 9pt

    # Salvar o documento modificado
    doc.save(output_path)
    print(f"Arquivo salvo com sucesso em {output_path}")

def adicionar_nova_linha_fabricacao(caminho_arquivo, texto_inicial_tabela, novo_codigo, descricao, dcb, quantidade, formula, formula_unitaria, caminho_arquivo_salvo):
    # Carregar o documento Word
    documento = Document(caminho_arquivo)

    # Iterar pelas tabelas no documento
    for tabela in documento.tables:
        # Verificar se a tabela começa com o texto inicial desejado
        if tabela.rows[0].cells[0].text.strip().startswith(texto_inicial_tabela):
            # Criar uma nova linha no final da tabela
            nova_linha = tabela.add_row()

            # Preencher os dados da nova linha
            nova_linha.cells[0].text = ""
            nova_linha.cells[1].text = novo_codigo  # Coluna vazia
            nova_linha.cells[2].text = descricao
            nova_linha.cells[3].text = dcb
            nova_linha.cells[4].text = quantidade
            nova_linha.cells[5].text = formula
            nova_linha.cells[6].text = formula_unitaria

            # Reordenar a tabela: Mover a nova linha para a posição desejada
            nova_linha_xml = nova_linha._element  # Pegar o elemento XML da nova linha
            tabela._tbl.remove(nova_linha_xml)  # Remover do final
            tabela._tbl.insert(4, nova_linha_xml)  # Inserir sem substituir nenhuma linha (índice 5)

            break  # Processar apenas a primeira tabela correspondente

    # Salvar o documento atualizado
    documento.save(caminho_arquivo_salvo)